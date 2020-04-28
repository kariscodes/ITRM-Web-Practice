import k_search

#Python3: install python-docx package (but the package name is the same as 'docx'
from datetime import datetime

import docx
from docx import Document
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def do_write_each(web_driver, keywords):
    date = datetime.today()
    dateStr = date.strftime('%Y%m%d')

    for input_keyword in keywords:
        s = k_search.Search()
        try:
            titles, links = s.search_in_goolenews(web_driver, input_keyword, 1)
        except Exception as e:
            print(e)
        document = Document()
        # Document Title
        document.add_heading('오늘의 검색 결과', 0)
        document.add_paragraph('Keyword: ' + input_keyword)

        # Print a series of lists together
        for title, link in zip(titles, links):
            # document.add_paragraph(title, style='Quote')
            # document.add_paragraph(link)
            p = document.add_paragraph()
            add_hyperlink(p, title, link)

        # Save the each file
        filename = dateStr + '_' + input_keyword + '.docx'
        document.save(filename)

def do_write_all(web_driver, keywords):
    date = datetime.today()
    dateStr = date.strftime('%Y%m%d')

    document = Document()
    p_format = document.styles['Normal'].paragraph_format
    p_format.space_before = Pt(20)
    document.add_heading('오늘의 검색 결과', 0)
    for input_keyword in keywords:
        s = k_search.Search()
        try:
            titles, links = s.search_in_goolenews(web_driver, input_keyword, 1)
        except Exception as e:
            print(e)
        document.add_paragraph('Keyword: ' + input_keyword, style='Heading 2')
        # Print a series of lists together
        for title, link in zip(titles, links):
            # document.add_paragraph(title, style='Quote')
            # document.add_paragraph(link)
            p = document.add_paragraph()
            add_hyperlink(p, title, link)

    # Save the one file
    filename = dateStr + '_' + 'search_results' + '.docx'
    document.save(filename)